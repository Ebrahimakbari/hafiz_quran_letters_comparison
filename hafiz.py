import re
from collections import Counter
from docx import Document
import matplotlib.pyplot as plt



def normalize_letters(text):
    text = re.sub(r'[ئإأؤء]', 'ء', text)
    text = re.sub(r'[يى]', 'ی', text)
    text = re.sub(r'[هة]', 'ه', text)
    text = re.sub(r'[ٰٓٱ]', 'ا', text)
    text = text.replace("آ", "ا")
    text = text.replace("ك", "ک")         
    return text


def read_docx(file_path, fraction=1.0):
    doc = Document(file_path)
    full_text = ""
    total_paragraphs = len(doc.paragraphs)
    paragraphs_to_read = int(total_paragraphs * fraction)
    
    for para in doc.paragraphs[:paragraphs_to_read]:
        full_text += para.text + "\n"
    return full_text


def calculate_letter_probabilities(file_path, fraction=1.0):
    text = read_docx(file_path, fraction=fraction)
    # حذف کاراکترهای غیر از حروف فارسی
    persian_letters_pattern = r'[^آ-ی]'
    clean_text = re.sub(persian_letters_pattern, '', text)
    # جایگزینی حروف هم‌ارز
    normalized_text = normalize_letters(clean_text)
    # شمارش تعداد تکرار هر حرف
    letter_frequency = Counter(normalized_text)
    # محاسبه احتمال نسبی هر حرف
    total_count = sum(letter_frequency.values())
    letter_probabilities = {letter: count / total_count for letter,
                            count in letter_frequency.items()}
    return letter_probabilities


def plot_probability_density_comparison(file_path, fractions, output_image_path):
    # ترتیب حروف الفبای فارسی
    persian_alphabet = list("ابپتثجچحخدذرزژسشصضطظعغفقکگلمنوهی")
    plt.figure(figsize=(12, 6))
    colors = ['blue', 'green', 'red']  # رنگ‌های متفاوت برای هر مقدار fraction
    markers = ['o', 's', '^']          # نشانه‌های متفاوت برای هر مقدار fraction
    for i, fraction in enumerate(fractions):
        # محاسبه توزیع احتمال برای هر fraction
        letter_probabilities = calculate_letter_probabilities(file_path,
                                                                fraction=fraction)
        # مرتب‌سازی احتمالات بر اساس ترتیب الفبای فارسی
        sorted_probabilities = [letter_probabilities.get(letter, 0) for letter in persian_alphabet]
        # رسم نمودار برای این fraction
        plt.plot(persian_alphabet, sorted_probabilities, label=f'Fraction = {fraction}', 
                color=colors[i], marker=markers[i], linestyle='-', markersize=5)
        plt.fill_between(persian_alphabet, sorted_probabilities, color=colors[i], alpha=0.2)

    plt.xlabel("Letters")
    plt.ylabel("Relative probability")
    plt.title("Comparison based on different Fraction (hafiz)")
    plt.legend()
    plt.xticks(rotation=90)
    plt.savefig(output_image_path, format='png', dpi=300)
    plt.show()


file_path = 'hafiz.docx'
output_image_path = 'letter_probability_density_comparison.png'
fractions_to_compare = [1.0, 0.5, 0.25]
plot_probability_density_comparison(file_path, fractions_to_compare,
                                    output_image_path)
